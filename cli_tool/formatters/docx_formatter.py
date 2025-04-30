from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate(code, output, metadata):
    doc = Document()

    # Right-aligned header
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.add_run(f'{metadata["name"]}\n').bold = True
    para.add_run(f'{metadata["class"]}\n')
    para.add_run(f'{metadata["roll_number"]}\n')

    # Center-aligned experiment title
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run(f'EXPERIMENT {metadata["experiment_number"]}').bold = True

    # Code section
    doc.add_paragraph('Code:', style='Heading 2').bold = True
    doc.add_paragraph(code, style='Normal')

    # Output section
    doc.add_paragraph('Output:', style='Heading 2').bold = True   
    doc.add_paragraph(output, style='Normal')

    filename = f"experiment_{metadata['experiment_number']}.docx"
    doc.save(filename)
    print(f"Document generated: {filename}")
